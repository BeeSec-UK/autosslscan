#!/usr/bin/env python3
import argparse
import datetime
import os
import re
import socket
import subprocess
from multiprocessing import Pool

from libnmap.parser import NmapParser, NmapParserException
from lxml import etree as ET

COLOURS = {
    "blue": "\033[1;34m",
    "green": "\033[1;32m",
    "red": "\033[1;31m",
    "yellow": "\033[1;33m",
    "reset": "\033[0m"
}
SYMBOLS = {
    "plus": f"{COLOURS['blue']}[{COLOURS['reset']}{COLOURS['green']}+{COLOURS['reset']}{COLOURS['blue']}]",
    "minus": f"{COLOURS['blue']}[{COLOURS['reset']}{COLOURS['red']}-{COLOURS['reset']}{COLOURS['blue']}]",
    "cross": f"{COLOURS['blue']}[{COLOURS['reset']}{COLOURS['red']}x{COLOURS['reset']}{COLOURS['blue']}]",
    "star": f"{COLOURS['green']}[*]{COLOURS['reset']}{COLOURS['green']}",
    "warn": f"{COLOURS['blue']}[{COLOURS['reset']}{COLOURS['yellow']}!{COLOURS['reset']}{COLOURS['blue']}]",
    "end": f"{COLOURS['reset']}"
}

starttls_services_template = {
    "ftp": [],
    "imap": [],
    "irc": [],
    "ldap": [],
    "mysql": [],
    "pop3": [],
    "psql": [],
    "smtp": [],
    "xmpp": []
}


def banner():
    banner_text = f"""
    
    {COLOURS['yellow']}
              _                _                     
   __ _ _   _| |_ ___  ___ ___| |___  ___ __ _ _ __  
  / _` | | | | __/ _ \\/ __/ __| / __|/ __/ _` | '_ \\ 
 | (_| | |_| | || (_) \\__ \\__ \\ \\__ \\ (_| (_| | | | |
  \\__,_|\\__,_|\\__\\___/|___/___/_|___/\\___\\__,_|_| |_|
                                                     
    
    @BeeSec
    Helping you Bee Secure - https://github.com/BeeSec-UK/
    
    usage: autosslscan.py -d [directory-with-xml] -o [output-directory] -t [num-threads] -w [generate-word-doc]{COLOURS['reset']}
    
    """
    print(banner_text)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Auto-sslscan - SSL scanning for multiple Nmap XML reports.")
    parser.add_argument("-d", "--dir", dest="input_directory", required=True,
                        help="Path to the directory containing Nmap XML files")
    parser.add_argument("-o", "--output", dest="output_directory", required=True, help="Path to the output directory")
    parser.add_argument("-t", "--threads", dest="num_threads", type=int, default=10,
                        help="Number of threads for parallel execution")
    parser.add_argument("-v", "--verbose", action="store_true", help="Enable verbose output")
    parser.add_argument("-w", "--word", action="store_true", help="Generate Word document report")
    parser.add_argument("--timeout", type=int, default=30,
                        help="Timeout in seconds for each scan")
    return parser.parse_args()


def remove_ansi_escape_sequences(text: str) -> str:
    ansi_escape = re.compile(r'\x1b\[[0-9;]*[a-zA-Z]')
    return ansi_escape.sub('', text)


def check_ssl_support(host, port, verbose=False):
    """
    Quickly check if SSL/TLS is supported on the given host and port.
    Uses openssl s_client and stops after detecting 'BEGIN CERTIFICATE'.
    Includes a timeout to avoid hanging.
    """
    # Try to resolve if it's a hostname
    resolved_ip = resolve_hostname(host)
    scan_target = resolved_ip if resolved_ip != host else host

    # Use the original host as servername for SNI
    servername = host

    if verbose:
        print(f"{SYMBOLS['plus']} Checking SSL/TLS support for {scan_target}:{port}")

    bash_command = f"""timeout 3s openssl s_client -connect {scan_target}:{port} -servername {servername} </dev/null 2>/dev/null"""
    try:
        result = subprocess.run(bash_command, shell=True, text=True, capture_output=True)
        if verbose:
            print(f"{SYMBOLS['plus']} OpenSSL exit code: {result.returncode}")
            if result.stdout:
                print(f"{SYMBOLS['plus']} OpenSSL output received")

        # Check for various SSL/TLS indicators
        indicators = [
            "BEGIN CERTIFICATE",
            "SSL-Session:",
            "Server certificate",
            "Protocol  :",
        ]

        for indicator in indicators:
            if indicator in result.stdout:
                if verbose:
                    print(f"{SYMBOLS['plus']} SSL/TLS support detected ({indicator})")
                return True

        if verbose:
            print(f"{SYMBOLS['minus']} No SSL/TLS indicators found in response")
        return False

    except Exception as e:
        if verbose:
            print(f"{SYMBOLS['cross']} Error checking SSL/TLS support: {e}")
        return False


def perform_ssl_scan_tls_service(host: str, service_name: str) -> tuple:
    ip_or_host, port = host.split(':')
    # Try to resolve if it's a hostname
    resolved_ip = resolve_hostname(ip_or_host)
    scan_target = resolved_ip if resolved_ip != ip_or_host else ip_or_host

    try:
        print(f"{SYMBOLS['plus']} Performing sslscan {ip_or_host}:{port} with STARTTLS {service_name}")
        result = subprocess.run(
            ["sslscan", "--no-sigs", f"--starttls-{service_name}", f"{scan_target}:{port}"],
            capture_output=True, text=True, check=True
        )
        print(f"{SYMBOLS['star']}{COLOURS['green']} Finished scanning {ip_or_host}:{port}")
        return ip_or_host, port, result.stdout
    except subprocess.CalledProcessError as e:
        print(f"{SYMBOLS['cross']} Error running sslscan for {ip_or_host}:{port}: {e}")
        return ip_or_host, port, None


def perform_ssl_scan(host: str) -> tuple:
    ip_or_host, port = host.split(':')
    resolved_ip = resolve_hostname(ip_or_host)
    scan_target = resolved_ip if resolved_ip != ip_or_host else ip_or_host

    print(f"{SYMBOLS['plus']} Performing sslscan {ip_or_host}:{port}")
    try:
        # First check if sslscan exists
        try:
            subprocess.run(["sslscan", "--version"], capture_output=True, check=True)
        except FileNotFoundError:
            print(f"{SYMBOLS['cross']} sslscan not found. Please install sslscan.")
            return ip_or_host, port, None

        result = subprocess.run(
            ["sslscan", "--no-colour", f"{scan_target}:{port}"],
            capture_output=True, text=True, check=True
        )
        print(f"{SYMBOLS['star']} Finished scanning {ip_or_host}:{port}")
        return ip_or_host, port, result.stdout
    except subprocess.CalledProcessError as e:
        print(f"{SYMBOLS['cross']} Error running sslscan for {ip_or_host}:{port}: {e}")
        return ip_or_host, port, None


def process_results(results, output_folders, vuln_folder):
    for result in results:
        if not result:
            continue
        ip, port, scan_output = result
        if not scan_output:
            continue

        temp = remove_ansi_escape_sequences(scan_output).splitlines()
        if scan_output and len(temp) > 4:
            check_ssl_wildcard(scan_output, f"{vuln_folder}/SSL_Wildcard_Present.txt", ip, port, output_folders)
            check_signed_cert_rsa_keylength(scan_output, f"{vuln_folder}/Weak_Signed_Certificate_RSA_Keylength.txt", ip,
                                            port, output_folders)
            check_tls_fallback(scan_output, f"{vuln_folder}/No_TLS_Fallback_SCSV_Support.txt", ip, port, output_folders)
            check_legacy_protocols(scan_output, f"{vuln_folder}/Legacy_SSL_And_TLS_Protocols.txt", ip, port,
                                   output_folders)
            check_medium_strength_ciphers(scan_output, f"{vuln_folder}/Medium_Strength_Cipher_Suites.txt", ip, port,
                                          output_folders)
            check_null_ciphers(scan_output, f"{vuln_folder}/NULL_Ciphers.txt", ip, port, output_folders)
            check_dhe_ciphers(scan_output, f"{vuln_folder}/Diffie_Hellman Modulus_<2048-bits.txt", ip, port,
                              output_folders)
            check_untrusted_certificate(scan_output, f"{vuln_folder}/Untrusted_Certificates.txt", ip, port,
                                        output_folders)
            check_cbc_ciphers(scan_output, f"{vuln_folder}/CBC_Cipher_Suites.txt", ip, port, output_folders)
            check_rc4_ciphers(scan_output, f"{vuln_folder}/RC4_Cipher_Suites.txt", ip, port, output_folders)
            check_certificate_expiry(scan_output, f"{vuln_folder}/Non_Valid_Certificates.txt", ip, port, output_folders)
            check_tls_v1_3_disabled(scan_output, f"{vuln_folder}/TLSv1.3_Disabled.txt", ip, port, output_folders)
            check_sha1_hash(scan_output, f"{vuln_folder}/Weak_Hashing_Algorithm.txt", ip, port, output_folders)
            check_self_signed_certificate(scan_output, f"{vuln_folder}/Self_Signed_Certificate.txt", ip, port,
                                          output_folders)
            check_wrong_hostname(scan_output, f"{vuln_folder}/Wrong_Hostname.txt", ip, port, output_folders)
        elif len(temp) <= 4:
            with open(f"{output_folders['sslscan_folder']}/errors.txt", "a") as f:
                f.write(f"{ip}:{port}\n")


def update_xml_with_resolved_ips(xml_file: str, verbose: bool = False):
    """Update XML file by replacing IPs of hostnames with their current resolved IPs."""
    try:
        tree = ET.parse(xml_file)
        root = tree.getroot()
        modified = False

        # Process hostname elements
        for hostname in root.findall(".//hostname[@name]"):
            host_name = hostname.get('name')
            resolved_ip = resolve_hostname(host_name, verbose)
            if resolved_ip != host_name:  # If resolution was successful
                # Find parent host element and update its address
                host = hostname.getparent()
                if host is not None and host.get('address'):
                    old_ip = host.get('address')
                    host.set('address', resolved_ip)
                    modified = True
                    if verbose:
                        print(f"{SYMBOLS['plus']} Updated host IP from {old_ip} to {resolved_ip} for {host_name}")

        # Process hop elements
        for hop in root.findall(".//hop[@host]"):
            host_name = hop.get('host')
            resolved_ip = resolve_hostname(host_name, verbose)
            if resolved_ip != host_name:  # If resolution was successful
                old_ip = hop.get('ipaddr')
                hop.set('ipaddr', resolved_ip)
                modified = True
                if verbose:
                    print(f"{SYMBOLS['plus']} Updated hop IP from {old_ip} to {resolved_ip} for {host_name}")

        if modified:
            tree.write(xml_file)
            if verbose:
                print(f"{SYMBOLS['star']} Updated XML file with resolved IPs: {xml_file}")

    except ET.ParseError:
        print(f"{SYMBOLS['cross']} Failed to parse XML file: {xml_file}")
    except Exception as e:
        print(f"{SYMBOLS['cross']} Error processing XML file {xml_file}: {e}")


def process_file(xml_file: str, output_folders: dict, num_threads: int, verbose: bool):
    try:
        # Update XML with resolved IPs before processing
        update_xml_with_resolved_ips(xml_file, verbose)

        report = NmapParser.parse_fromfile(xml_file)
        process_report(report, output_folders, num_threads, verbose)
    except NmapParserException:
        with open(f"{output_folders['sslscan_folder']}/errors.log", "a") as f:
            f.write(f"Failed to parse: {xml_file}\n")
        print(f"{SYMBOLS['cross']} Error parsing file: {xml_file}")


def resolve_hostname(hostname, verbose=False):
    """Resolve hostname to IP address. Return the IP or the original hostname if resolution fails."""
    try:
        if re.match(r'^(?:[0-9]{1,3}\.){3}[0-9]{1,3}$', hostname):
            if verbose:
                print(f"{SYMBOLS['plus']} {hostname} is already an IP address")
            return hostname  # Already an IP address
        ip = socket.gethostbyname(hostname)
        if verbose:
            print(f"{SYMBOLS['plus']} Resolved {hostname} to {ip}")
        return ip
    except socket.gaierror:
        if verbose:
            print(f"{SYMBOLS['warn']} Could not resolve hostname: {hostname}")
        return hostname


def process_report(report, output_folders, num_threads, verbose=False):
    starttls_services = {service_name: [] for service_name in starttls_services_template}
    ssl_services = []

    # Default ports to always check
    default_ssl_ports = [443, 8443]

    for host in report.hosts:
        if verbose:
            print(f"\n{SYMBOLS['star']} Processing host: {host.address}")
            if host.hostnames:
                print(f"{SYMBOLS['plus']} Found hostnames: {', '.join(host.hostnames)}")

        # Get all possible hostnames and resolve them to IPs first
        addresses = []
        original_to_resolved = {}  # Map original addresses/hostnames to resolved IPs

        # Resolve hostnames
        if host.hostnames:
            for hostname in host.hostnames:
                resolved_ip = resolve_hostname(hostname, verbose)
                if resolved_ip not in addresses:
                    addresses.append(resolved_ip)
                    original_to_resolved[hostname] = resolved_ip

        # If no hostnames were resolved, check if the original address needs resolving
        if not addresses:
            resolved_original = resolve_hostname(host.address, verbose)
            if resolved_original != host.address:
                addresses.append(resolved_original)
                original_to_resolved[host.address] = resolved_original
            else:
                addresses.append(host.address)

        if verbose:
            # Show all open ports that will be checked
            open_ports = []
            scanned_ports = set()  # Keep track of ports we'll scan

            # Add ports from nmap scan
            for s in host.services:
                if s.state == "open":
                    service_name = f" ({s.service})" if s.service else ""
                    open_ports.append(f"{s.port}{service_name}")
                    scanned_ports.add(s.port)

            # Add default SSL ports if not already included
            for port in default_ssl_ports:
                if port not in scanned_ports:
                    open_ports.append(f"{port} (default check)")

            if open_ports:
                print(f"{SYMBOLS['plus']} Found open ports: {', '.join(map(str, open_ports))}")
                print(f"{SYMBOLS['plus']} Will check these ports on resolved IPs: {', '.join(addresses)}")
            else:
                print(f"{SYMBOLS['minus']} No open ports found, will check default SSL ports")

        # Use the first resolved IP or the original address if no resolution
        main_address = addresses[0] if addresses else host.address

        tls_services_found = False
        # First check ports from nmap scan
        for s in host.services:
            if s.state == "open":
                if verbose:
                    service_name = f" ({s.service})" if s.service else ""
                    print(f"{SYMBOLS['plus']} Checking port {s.port}{service_name} on {main_address}")

                if check_ssl_support(main_address, s.port, verbose):
                    ssl_services.append(f"{main_address}:{s.port}")
                    tls_services_found = True
                    if verbose:
                        print(f"{SYMBOLS['plus']} Found SSL/TLS service on {main_address}:{s.port}")
                elif s.service in starttls_services:
                    starttls_services[s.service].append(f"{main_address}:{s.port}")
                    tls_services_found = True
                    if verbose:
                        print(f"{SYMBOLS['plus']} Found STARTTLS service ({s.service}) on {main_address}:{s.port}")
                elif verbose:
                    print(f"{SYMBOLS['minus']} No SSL/TLS support detected on {main_address}:{s.port}")

        # Then check default SSL ports if they weren't in the nmap scan
        scanned_ports = {s.port for s in host.services if s.state == "open"}
        for port in default_ssl_ports:
            if port not in scanned_ports:
                if verbose:
                    print(f"{SYMBOLS['plus']} Checking default SSL port {port} on {main_address}")

                if check_ssl_support(main_address, port, verbose):
                    ssl_services.append(f"{main_address}:{port}")
                    tls_services_found = True
                    if verbose:
                        print(f"{SYMBOLS['plus']} Found SSL/TLS service on {main_address}:{port}")
                elif verbose:
                    print(f"{SYMBOLS['minus']} No SSL/TLS support detected on {main_address}:{port}")

        if verbose and not tls_services_found:
            print(f"{SYMBOLS['minus']} No SSL/TLS services found for this host")

    # Remove any duplicates that might have been created
    ssl_services = list(set(ssl_services))
    for service in starttls_services:
        starttls_services[service] = list(set(starttls_services[service]))

    with Pool(processes=num_threads) as pool:
        ssl_results = []
        if ssl_services:
            ssl_results = pool.map(perform_ssl_scan, ssl_services)

        starttls_results = []
        starttls_targets = [(host, service) for service, hosts in starttls_services.items() for host in hosts]
        if starttls_targets:
            starttls_results = pool.starmap(perform_ssl_scan_tls_service, starttls_targets)

    process_results(ssl_results, output_folders, output_folders["vuln_output"])
    process_results(starttls_results, output_folders, output_folders["starttls_vuln_output"])


def write_final_results(vuln_folder):
    with open(f"{vuln_folder}/Final_Results.txt", 'a') as f:
        for text_file in os.listdir(vuln_folder):
            if text_file != 'Final_Results.txt':
                title = f"{text_file.replace('_', ' ').replace('.txt', '')}"
                with open(f"{vuln_folder}/{text_file}", 'r') as s:
                    results = s.read()
                    if results.strip():
                        f.write(f"{title}:\n{results}\n")


def consolidate_results(vuln_folder, sslscan_folder, label):
    write_final_results(vuln_folder)
    final_file = os.path.join(sslscan_folder, "Big_Final_Results.txt")
    with open(final_file, "a") as big_final:
        big_final.write(f"\n{label} Services:\n")
        with open(os.path.join(vuln_folder, "Final_Results.txt"), 'r') as f:
            contents = f.read()
            if contents.strip():
                big_final.write(contents)


##############################
# XML Consolidation
##############################

# Requested Category Mapping
# Map the existing txt files to the requested categories.
category_mapping = {
    "Legacy_SSL_And_TLS_Protocols.txt": "Legacy Protocols (SSLv2/3, TLSv1.0/1.1)",
    "Non_Valid_Certificates.txt": "SSL Certificate Expired",
    "Cert_Expires_Within_6_Months.txt": "SSL Certificate Expires within 6 months",
    "Weak_Hashing_Algorithm.txt": "SSL Certificate Signed Using Weak Hashing Algorithm",
    "Weak_Signed_Certificate_RSA_Keylength.txt": "SSL Certificate Chain Contains RSA Keys Less Than 2048 bits",
    "CBC_Cipher_Suites.txt": "SSL Cipher Block Chaining Cipher Suites Supported",
    "Medium_Strength_Cipher_Suites.txt": "SSL Medium Strength Cipher Suites Supported (SWEET32)",
    "Untrusted_Certificates.txt": "SSL Certificate Cannot be Trusted",
    "Wrong_Hostname.txt": "SSL Certificate with Wrong Hostname",
    "Self_Signed_Certificate.txt": "Self-Signed Certificate",
    "RC4_Cipher_Suites.txt": "SSL RC4 Cipher Suites Supported (Bar Mitsvah)",
    "NULL_Ciphers.txt": "NULL Cipher Support",
    "Diffie_Hellman Modulus_<2048-bits.txt": "Weak DH Parameters",
    "No_TLS_Fallback_SCSV_Support.txt": "No TLS Fallback Support",
    "SSL_Wildcard_Present.txt": "Wildcard Certificate Present",
    "TLSv1.3_Disabled.txt": "TLSv1.3 Not Supported"
}


def parse_ip_port(line):
    line = line.strip()
    if ":" in line:
        # Split only on the first colon to preserve additional information
        parts = line.split(":", 1)
        ip = parts[0].strip()
        # Keep all the port information including any parenthetical details
        port = parts[1].strip()
        return ip, port
    return None, None


def consolidate_xml_results(vuln_output_folder, starttls_vuln_folder, sslscan_folder):
    vulnerabilities = {}

    def add_finding(category, ip, port):
        if category not in vulnerabilities:
            vulnerabilities[category] = {}
        if ip not in vulnerabilities[category]:
            vulnerabilities[category][ip] = set()
        # Extract just the port number from the port string
        port_number = port.split()[0].split('(')[0].strip()
        vulnerabilities[category][ip].add(port)

    # Process files in ssl vuln folder
    for fname in os.listdir(vuln_output_folder):
        if fname == "Final_Results.txt":
            continue
        category = category_mapping.get(fname)
        if category:
            with open(os.path.join(vuln_output_folder, fname), 'r') as f:
                for line in f:
                    ip, port = parse_ip_port(line)
                    if ip and port:
                        add_finding(category, ip, port)

    # Process files in starttls vuln folder
    for fname in os.listdir(starttls_vuln_folder):
        if fname == "Final_Results.txt":
            continue
        category = category_mapping.get(fname)
        if category:
            with open(os.path.join(starttls_vuln_folder, fname), 'r') as f:
                for line in f:
                    ip, port = parse_ip_port(line)
                    if ip and port:
                        add_finding(category, ip, port)

    # Build the XML
    root = ET.Element("vulnerabilities")

    def port_sort_key(port_str):
        """Helper function to sort ports properly"""
        # Extract the port number from the string (everything before space or parenthesis)
        port_number = port_str.split()[0].split('(')[0].strip()
        try:
            return (0, int(port_number))  # Numbers come first
        except ValueError:
            return (1, port_str)  # Strings come second

    for category, hosts in vulnerabilities.items():
        vuln_el = ET.SubElement(root, "vulnerability", attrib={"name": category})
        for host, ports in hosts.items():
            host_el = ET.SubElement(vuln_el, "ip_or_hostname", attrib={"name": host})
            # Use the new port_sort_key function
            for p in sorted(ports, key=port_sort_key):
                port_el = ET.SubElement(host_el, "affected_port")
                port_el.text = p

    output_xml = os.path.join(sslscan_folder, "Big_Final_Results.xml")
    tree = ET.ElementTree(root)
    tree.write(output_xml, encoding="utf-8", xml_declaration=True)
    print(f"{SYMBOLS['plus']} XML results consolidated in {output_xml}")

    return vulnerabilities


##############################
# Vulnerability Check Functions
##############################

def check_legacy_protocols(scan_output: str, result_path: str, host: str, port: str,
                           output_folders: dict[str, str]) -> None:
    legacy_protocols = ["SSLv2", "SSLv3", "TLSv1.0", "TLSv1.1"]
    scan_output = remove_ansi_escape_sequences(scan_output)
    with open(f'{output_folders["raw_output"]}/{host}:{port}.txt', 'w') as f:
        f.write(scan_output)

    enabled_protocols = []
    for line in scan_output.splitlines():
        parts = line.split()
        if len(parts) == 2 and parts[1].lower() == "enabled" and parts[0] in legacy_protocols:
            enabled_protocols.append(parts[0])

    if enabled_protocols:
        with open(result_path, 'a') as f:
            protocols_str = ", ".join(enabled_protocols)
            f.write(f"{host}:{port} (Enabled: {protocols_str})\n")


def check_tls_v1_3_disabled(scan_output: str, result_path: str, host: str, port: str,
                            output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        parts = line.split()
        if len(parts) == 2 and parts[1].lower() == "disabled" and parts[0] == "TLSv1.3":
            with open(result_path, 'a') as s:
                s.write(f"{host}:{port}\n")
            return


def check_certificate_expiry(scan_output: str, result_path: str, host: str, port: str,
                             output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    valid_from = None
    valid_until = None
    subject = None
    issuer = None

    for line in scan_output.splitlines():
        if line.startswith("Not valid before:"):
            valid_from = line.replace("Not valid before:", "").strip()
        elif line.startswith("Not valid after:"):
            valid_until = line.replace("Not valid after:", "").strip()
        elif line.startswith("Subject:"):
            subject = line.replace("Subject:", "").strip()
        elif line.startswith("Issuer:"):
            issuer = line.replace("Issuer:", "").strip()

    if valid_from and valid_until:
        current_date = datetime.datetime.now(datetime.UTC)
        valid_from_date = datetime.datetime.strptime(valid_from, "%b %d %H:%M:%S %Y %Z").replace(tzinfo=datetime.UTC)
        valid_until_date = datetime.datetime.strptime(valid_until, "%b %d %H:%M:%S %Y %Z").replace(tzinfo=datetime.UTC)
        days_remaining = (valid_until_date - current_date).days

        # Check if certificate is expired or not yet valid
        if current_date < valid_from_date or current_date > valid_until_date:
            with open(result_path, 'a') as s:
                s.write(f"{host}:{port} (Certificate expired or not yet valid, Subject: {subject})\n")
        # Check if certificate expires within 6 months (approximately 182 days)
        elif days_remaining <= 182:
            with open(os.path.join(output_folders["vuln_output"], "Cert_Expires_Within_6_Months.txt"), 'a') as s:
                s.write(f"{host}:{port} (Expires in {days_remaining} days, Subject: {subject})\n")


def check_signed_cert_rsa_keylength(scan_output: str, result_path: str, host: str, port: str,
                                    output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        if line.startswith('RSA Key Strength:'):
            parts = line.split()
            if int(parts[3]) < 2048:
                with open(result_path, 'a') as s:
                    s.write(f"{host}:{port}\n")
                return


def check_tls_fallback(scan_output: str, result_path: str, host: str, port: str,
                       output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        if line.strip() == 'Server does not support TLS Fallback SCSV':
            with open(result_path, 'a') as s:
                s.write(f"{host}:{port} (TLS Fallback SCSV not supported)\n")
            return


def check_dhe_ciphers(scan_output: str, result_path: str, host: str, port: str, output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        parts = line.split()
        if len(parts) >= 7:
            if parts[4].startswith("DHE") and not any("Curve" in part for part in parts):
                # Adjust threshold if needed for strict 1024-bit check
                if int(parts[6]) < 2024:
                    with open(result_path, 'a') as s:
                        s.write(f"{host}:{port}\n")
                    return


def check_untrusted_certificate(scan_output: str, result_path: str, host: str, port: str,
                                output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        # If the Issuer line is printed in red by sslscan (indicating untrusted),
        # you may detect it with certain patterns. Adjust as needed.
        if "Issuer:" in line and "\x1b[31m" in line:
            with open(result_path, 'a') as s:
                s.write(f"{host}:{port}\n")
            return


def check_cbc_ciphers(scan_output: str, result_path: str, host: str, port: str, output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        if "CBC" in line:
            with open(result_path, 'a') as f:
                f.write(f"{host}:{port}\n")
            return


def check_sha1_hash(scan_output: str, result_path: str, host: str, port: str, output_folders: dict[str, str]) -> None:
    # Also detect MD5 signatures as weak hashing algorithm
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        if "SHA-1" in line.lower() or "md5withrsaencryption" in line.lower():
            with open(result_path, 'a') as f:
                f.write(f"{host}:{port}\n")
            return


def check_rc4_ciphers(scan_output: str, result_path: str, host: str, port: str, output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        if "RC4" in line:
            with open(result_path, 'a') as f:
                f.write(f"{host}:{port}\n")
            return


def check_medium_strength_ciphers(scan_output: str, result_path: str, host: str, port: str,
                                  output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        if not line.startswith("OpenSSL"):
            parts = line.split()
            if (len(parts) >= 3 and parts[2].isdigit() and 1 <= int(parts[2]) < 128):
                with open(result_path, 'a') as f:
                    f.write(f"{host}:{port}\n")
                return


def check_null_ciphers(scan_output: str, result_path: str, host: str, port: str,
                       output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    for line in scan_output.splitlines():
        if "NULL" in line:
            with open(result_path, 'a') as f:
                f.write(f"{host}:{port}\n")
            return


def check_ssl_wildcard(scan_output: str, result_path: str, host: str, port: str,
                       output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    wildcard_found = False
    cn_value = None

    for line in scan_output.splitlines():
        if line.startswith("Subject:"):
            cn_match = re.search(r"CN\s*=\s*([^,/]+)", line)
            if cn_match:
                cn_value = cn_match.group(1)
                if '*' in cn_value:
                    wildcard_found = True
                    break
        elif "DNS:" in line:
            dns_names = re.findall(r"DNS:([^,\s]+)", line)
            for name in dns_names:
                if '*' in name:
                    cn_value = name
                    wildcard_found = True
                    break

    if wildcard_found and cn_value:
        with open(result_path, 'a') as f:
            f.write(f"{host}:{port} (Wildcard: {cn_value})\n")


def check_self_signed_certificate(scan_output: str, result_path: str, host: str, port: str,
                                  output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    subject = None
    issuer = None

    for line in scan_output.splitlines():
        if line.startswith("Subject:"):
            subject = line.replace("Subject:", "").strip()
        elif line.startswith("Issuer:"):
            issuer = line.replace("Issuer:", "").strip()

    # If subject and issuer are identical, it's a self-signed certificate
    if subject and issuer and subject == issuer:
        with open(result_path, 'a') as s:
            s.write(f"{host}:{port} (Self-signed certificate)\n")


def check_wrong_hostname(scan_output: str, result_path: str, host: str, port: str,
                         output_folders: dict[str, str]) -> None:
    scan_output = remove_ansi_escape_sequences(scan_output)
    hostname = host

    # Try to extract domain from hostname if it's an IP
    if re.match(r'^(?:[0-9]{1,3}\.){3}[0-9]{1,3}$', hostname):
        hostname = None  # Skip check for IPs unless we can determine expected hostname

    if hostname:
        cn_match = False
        san_match = False
        common_name = None

        for line in scan_output.splitlines():
            if "CN=" in line:
                cn_match = re.search(r"CN\s*=\s*([^,/]+)", line)
                if cn_match:
                    common_name = cn_match.group(1).strip()
                    if common_name.startswith("*"):
                        # Check wildcard domain
                        domain_part = common_name[1:]  # Remove the *
                        if hostname.endswith(domain_part):
                            cn_match = True
                            break
                    # Direct match
                    elif common_name.lower() == hostname.lower():
                        cn_match = True
                        break

            # Check Subject Alternative Names
            if "DNS:" in line:
                dns_names = re.findall(r"DNS:([^,\s]+)", line)
                for name in dns_names:
                    if name.startswith("*"):
                        # Check wildcard domain
                        domain_part = name[1:]  # Remove the *
                        if hostname.endswith(domain_part):
                            san_match = True
                            break
                    elif name.lower() == hostname.lower():
                        san_match = True
                        break

        if not (cn_match or san_match):
            with open(result_path, 'a') as s:
                s.write(f"{host}:{port} (Certificate hostname mismatch, CN={common_name or 'unknown'})\n")


def print_summary_statistics(output_folders: dict):
    print(f"\n{SYMBOLS['star']} Scan Summary:")

    total_hosts = set()
    findings = {}

    for category_file in os.listdir(output_folders["vuln_output"]):
        if category_file == "Final_Results.txt":
            continue

        category = category_mapping.get(category_file, category_file)
        file_path = os.path.join(output_folders["vuln_output"], category_file)

        with open(file_path) as f:
            lines = f.readlines()
            if lines:
                findings[category] = len(lines)
                for line in lines:
                    host = line.split(":")[0]
                    total_hosts.add(host)

    print(f"\nTotal unique hosts with findings: {len(total_hosts)}")
    print("\nFindings by category:")
    for category, count in findings.items():
        print(f"  {category}: {count}")


def generate_word_report(output_folders, vulnerabilities):
    try:
        from docx import Document
        from docx.shared import Inches

        document = Document()
        document.add_heading('SSL/TLS Scan Report', 0)

        # Add timestamp
        document.add_paragraph(f'Report generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        # Add summary
        document.add_heading('Summary of Findings', level=1)
        summary_table = document.add_table(rows=1, cols=2)
        summary_table.style = 'Table Grid'
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = 'Issue'
        header_cells[1].text = 'Count'

        total_issues = 0

        # Map file names to issue titles
        issue_titles = {
            "Legacy_SSL_And_TLS_Protocols.txt": "Legacy Protocols (SSLv2/3, TLSv1.0/1.1)",
            "Non_Valid_Certificates.txt": "SSL Certificate Expired",
            "Cert_Expires_Within_6_Months.txt": "SSL Certificate Expires within 6 months",
            "Weak_Hashing_Algorithm.txt": "SSL Certificate Signed Using Weak Hashing Algorithm",
            "Weak_Signed_Certificate_RSA_Keylength.txt": "SSL Certificate Chain Contains RSA Keys Less Than 2048 bits",
            "CBC_Cipher_Suites.txt": "SSL Cipher Block Chaining Cipher Suites Supported",
            "Medium_Strength_Cipher_Suites.txt": "SSL Medium Strength Cipher Suites Supported (SWEET32)",
            "Untrusted_Certificates.txt": "SSL Certificate Cannot be Trusted",
            "Wrong_Hostname.txt": "SSL Certificate with Wrong Hostname",
            "Self_Signed_Certificate.txt": "Self-Signed Certificate",
            "RC4_Cipher_Suites.txt": "SSL RC4 Cipher Suites Supported (Bar Mitsvah)"
        }

        issue_data = {}

        # Get all files in vulnerability folders
        vuln_files = []
        for file in os.listdir(output_folders["vuln_output"]):
            if file.endswith('.txt') and file != 'Final_Results.txt':
                vuln_files.append((file, os.path.join(output_folders["vuln_output"], file)))

        for file in os.listdir(output_folders["starttls_vuln_output"]):
            if file.endswith('.txt') and file != 'Final_Results.txt':
                vuln_files.append((file, os.path.join(output_folders["starttls_vuln_output"], file)))

        # Get count and data for each issue
        for file_name, file_path in vuln_files:
            if file_name in issue_titles:
                issue_title = issue_titles[file_name]
                hosts = []

                with open(file_path, 'r') as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            hosts.append(line)

                count = len(hosts)
                if count > 0:
                    total_issues += count
                    issue_data[issue_title] = hosts

                    # Add to summary table
                    row_cells = summary_table.add_row().cells
                    row_cells[0].text = issue_title
                    row_cells[1].text = str(count)

        document.add_paragraph(f'Total Issues Found: {total_issues}')

        # Add detailed findings
        document.add_heading('Detailed Findings', level=1)

        for issue_title, hosts in issue_data.items():
            document.add_heading(issue_title, level=2)

            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Issue'
            header_cells[1].text = 'Affected Hosts'

            row_cells = table.add_row().cells
            row_cells[0].text = issue_title

            # Format all hosts into one cell
            hosts_text = '\n'.join(hosts)
            row_cells[1].text = hosts_text

        # Save the document
        document.save(os.path.join(output_folders["sslscan_folder"], 'SSL_Scan_Report.docx'))
        print(
            f"{SYMBOLS['plus']} Word report generated: {os.path.join(output_folders['sslscan_folder'], 'SSL_Scan_Report.docx')}")

    except ImportError:
        print(
            f"{SYMBOLS['cross']} Could not generate Word document. Please install python-docx package: pip install python-docx")


def main():
    args = parse_args()
    input_directory = args.input_directory
    output_directory = args.output_directory
    num_threads = args.num_threads
    verbose = args.verbose
    generate_word = args.word

    sslscan_folder = os.path.join(output_directory, "sslscan")
    os.makedirs(sslscan_folder, exist_ok=True)

    raw_output_folder = os.path.join(sslscan_folder, "raw_output")
    os.makedirs(raw_output_folder, exist_ok=True)

    vuln_output_folder = os.path.join(sslscan_folder, "ssl")
    os.makedirs(vuln_output_folder, exist_ok=True)

    starttls_vuln_folder = os.path.join(sslscan_folder, "starttls")
    os.makedirs(starttls_vuln_folder, exist_ok=True)

    output_folders = {
        "raw_output": raw_output_folder,
        "vuln_output": vuln_output_folder,
        "starttls_vuln_output": starttls_vuln_folder,
        "sslscan_folder": sslscan_folder
    }

    with open(f"{sslscan_folder}/errors.log", "w") as f:
        pass
    with open(os.path.join(sslscan_folder, "Big_Final_Results.txt"), "w") as f:
        pass

    files_to_make = [
        "Legacy_SSL_And_TLS_Protocols.txt",
        "Non_Valid_Certificates.txt",
        "Cert_Expires_Within_6_Months.txt",
        "NULL_Ciphers.txt",
        "Diffie_Hellman Modulus_<2048-bits.txt",
        "Untrusted_Certificates.txt",
        "Medium_Strength_Cipher_Suites.txt",
        "CBC_Cipher_Suites.txt",
        "RC4_Cipher_Suites.txt",
        "No_TLS_Fallback_SCSV_Support.txt",
        "Weak_Signed_Certificate_RSA_Keylength.txt",
        "Weak_Hashing_Algorithm.txt",
        "SSL_Wildcard_Present.txt",
        "TLSv1.3_Disabled.txt",
        "Self_Signed_Certificate.txt",
        "Wrong_Hostname.txt",
        "Final_Results.txt"
    ]

    for file in files_to_make:
        open(os.path.join(vuln_output_folder, file), 'w').close()
        open(os.path.join(starttls_vuln_folder, file), 'w').close()

    xml_files = [os.path.join(input_directory, f) for f in os.listdir(input_directory) if f.endswith(".xml")]

    banner()

    for xml_file in xml_files:
        print(f"{SYMBOLS['plus']} Processing file: {xml_file}")
        process_file(xml_file, output_folders, num_threads, verbose)

    # Consolidate all results
    consolidate_results(output_folders["vuln_output"], output_folders["sslscan_folder"], "SSL")
    consolidate_results(output_folders["starttls_vuln_output"], output_folders["sslscan_folder"], "STARTTLS")

    # Produce XML results
    vulnerabilities = consolidate_xml_results(output_folders["vuln_output"], output_folders["starttls_vuln_output"],
                                              output_folders["sslscan_folder"])

    # Generate Word document if requested
    if generate_word:
        generate_word_report(output_folders, vulnerabilities)

    print(
        f"{SYMBOLS['plus']} Results consolidated in {os.path.join(output_folders['sslscan_folder'], 'Big_Final_Results.txt')}")
    print_summary_statistics(output_folders)
    banner()


if __name__ == "__main__":
    main()
